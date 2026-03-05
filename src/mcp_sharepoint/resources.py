import base64, os, fitz, io, logging, time, pandas as pd
from typing import Dict, Any, List, Optional
from docx import Document
from openpyxl import load_workbook
from .common import logger, SHP_DOC_LIBRARY, sp_context, retry_on_connection_error

logger = logging.getLogger(__name__)

# Configuration
FILE_TYPES = {
    'text': ['.txt', '.csv', '.json', '.xml', '.html', '.md', '.js', '.css', '.py'],
    'pdf': ['.pdf'],
    'excel': ['.xlsx', '.xls'],
    'word': ['.docx', '.doc']
}

# Tree configuration from environment variables with defaults
TREE_CONFIG = {
    'max_depth': int(os.getenv('SHP_MAX_DEPTH', '15')),
    'max_folders_per_level': int(os.getenv('SHP_MAX_FOLDERS_PER_LEVEL', '100')),
    'level_delay': float(os.getenv('SHP_LEVEL_DELAY', '0.5'))
}

# Download configuration
DOWNLOAD_CONFIG = {
    'fallback_dir': './downloads'
}

def _get_sp_path(sub_path: Optional[str] = None) -> str:
    """Create a properly formatted SharePoint path"""
    return f"{SHP_DOC_LIBRARY}/{sub_path or ''}".rstrip('/')

def _ensure_directory_exists(directory: str) -> bool:
    """Ensure target directory exists, create if necessary"""
    try:
        os.makedirs(directory, exist_ok=True)
        return True
    except Exception as e:
        logger.error(f"Failed to create directory {directory}: {e}")
        return False

def _get_fallback_path(file_name: str) -> str:
    """Generate fallback path for downloads"""
    fallback_dir = DOWNLOAD_CONFIG['fallback_dir']
    _ensure_directory_exists(fallback_dir)
    return os.path.join(fallback_dir, file_name)

def _save_content_to_file(content_bytes: bytes, file_path: str) -> Dict[str, Any]:
    """Save binary content to local file with error handling"""
    try:
        # Ensure directory exists
        directory = os.path.dirname(file_path)
        if directory and not _ensure_directory_exists(directory):
            raise Exception(f"Cannot create directory: {directory}")
            
        # Write file
        with open(file_path, 'wb') as f:
            f.write(content_bytes)
        
        # Verify file was created
        if os.path.exists(file_path) and os.path.getsize(file_path) == len(content_bytes):
            return {"success": True, "path": os.path.abspath(file_path), "size": len(content_bytes)}
        else:
            raise Exception("File verification failed")
            
    except Exception as e:
        logger.error(f"Failed to save file to {file_path}: {e}")
        return {"success": False, "error": str(e)}

def _load_sp_items(path: str, item_type: str) -> List[Dict[str, Any]]:
    """Generic function to load folders or files from SharePoint"""
    @retry_on_connection_error(max_retries=3, delay=1.0)
    def _execute():
        folder = sp_context.web.get_folder_by_server_relative_url(path)
        items = getattr(folder, item_type)
        props = ["ServerRelativeUrl", "Name", "TimeCreated", "TimeLastModified"] + (["Length"] if item_type == "files" else [])
        sp_context.load(items, props)
        sp_context.execute_query()
        
        return [{
            "name": item.name,
            "url": item.properties.get("ServerRelativeUrl"),
            **({"size": item.properties.get("Length")} if item_type == "files" else {}),
            "created": item.properties.get("TimeCreated").isoformat() if item.properties.get("TimeCreated") else None,
            "modified": item.properties.get("TimeLastModified").isoformat() if item.properties.get("TimeLastModified") else None
        } for item in items]
    
    return _execute()

def list_folders(parent_folder: Optional[str] = None) -> List[Dict[str, Any]]:
    """List folders in the specified directory or root if not specified"""
    logger.info(f"Listing folders in {parent_folder or 'root directory'}")
    return _load_sp_items(_get_sp_path(parent_folder), "folders")

def list_documents(folder_name: str) -> List[Dict[str, Any]]:
    """List all documents in a specified folder"""
    logger.info(f"Listing documents in folder: {folder_name}")
    return _load_sp_items(_get_sp_path(folder_name), "files")

def extract_text_from_pdf(pdf_content):
    """Extract text from PDF using PyMuPDF"""
    try:
        pdf_document = fitz.open(stream=pdf_content, filetype="pdf")
        text_content = "".join(pdf_document[i].get_text() + "\n" for i in range(len(pdf_document)))
        page_count = len(pdf_document)
        pdf_document.close()
        return text_content.strip(), page_count
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {e}")
        raise

def extract_text_from_excel(content_bytes):
    """Extract text from Excel files"""
    try:
        sheets = pd.read_excel(io.BytesIO(content_bytes), sheet_name=None)
        text_parts = []
        for sheet_name, df in sheets.items():
            text_parts.append(f"=== {sheet_name} ===")
            text_parts.extend(df.head(50).fillna('').astype(str).apply(' | '.join, axis=1).tolist())
        return "\n".join(text_parts), len(sheets)
    except Exception as e:
        logger.error(f"Error extracting text from Excel: {e}")
        raise

def extract_text_from_word(content_bytes):
    """Extract text from Word documents"""
    try:
        doc = Document(io.BytesIO(content_bytes))
        text_parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                text_parts.append(" | ".join(cell.text.strip() for cell in row.cells))
        return "\n".join(text_parts), len(doc.paragraphs)
    except Exception as e:
        logger.error(f"Error extracting text from Word: {e}")
        raise

def get_folder_tree(parent_folder: Optional[str] = None) -> Dict[str, Any]:
    """Iteratively build folder tree level by level to avoid recursion limits"""
    root_path, tree_nodes = _get_sp_path(parent_folder), {}
    logger.info(f"Building iterative tree for {parent_folder or 'root'}")
    
    try:
        # Get root folder with retry logic
        @retry_on_connection_error(max_retries=3, delay=1.0)
        def _get_root():
            root = sp_context.web.get_folder_by_server_relative_url(root_path)
            sp_context.load(root, ["Name", "ServerRelativeUrl", "TimeCreated", "TimeLastModified"])
            sp_context.execute_query()
            return root
        
        root = _get_root()
        
        # Process folders level by level
        pending = [parent_folder or ""]
        for level in range(TREE_CONFIG['max_depth']):
            if not pending: break
            logger.info(f"Level {level + 1}: {len(pending)} folders")
            
            # Process all folders in this level by batches
            current_level_folders = pending.copy()
            next_level_folders = []
            pending = []  # Reset for next level
            
            # Process current level in batches to handle large numbers of folders
            while current_level_folders:
                batch = current_level_folders[:TREE_CONFIG['max_folders_per_level']]
                current_level_folders = current_level_folders[TREE_CONFIG['max_folders_per_level']:]
                
                for folder_path in batch:
                    try:
                        subfolders = [f["name"] for f in list_folders(folder_path)]
                        files = list_documents(folder_path)
                        
                        tree_nodes[folder_path] = [
                            {"name": name, "type": "folder", "children": []} for name in subfolders
                        ] + [{"name": f["name"], "path": f["url"], "type": "file", 
                             **{k: v for k, v in f.items() if k not in ["name", "url"]}} for f in files]
                        
                        # Add subfolders to next level processing
                        next_level_folders.extend([f"{folder_path}/{name}".strip('/') for name in subfolders])
                    except: 
                        logger.warning(f"Failed to process: {folder_path}")
                
                # Small delay between batches to avoid overwhelming SharePoint
                if current_level_folders:  # Only delay if more batches remain
                    time.sleep(0.1)
            
            # Set up for next level
            pending = next_level_folders
            
            if level < TREE_CONFIG['max_depth'] - 1: time.sleep(TREE_CONFIG['level_delay'])
        
        # Build nested structure
        def build_node(path: str) -> List[Dict]:
            children = tree_nodes.get(path, [])
            for child in children:
                if child["type"] == "folder":
                    child["children"] = build_node(f"{path}/{child['name']}".strip('/'))
            return children
        
        return {
            "name": root.name, "path": root.properties.get("ServerRelativeUrl"), "type": "folder",
            "created": root.properties.get("TimeCreated").isoformat() if root.properties.get("TimeCreated") else None,
            "modified": root.properties.get("TimeLastModified").isoformat() if root.properties.get("TimeLastModified") else None,
            "children": build_node(parent_folder or "")
        }
        
    except Exception as e:
        logger.error(f"Failed to build tree for '{root_path}': {e}")
        return {"name": os.path.basename(root_path), "path": root_path, "type": "folder", "error": "Could not access folder", "children": []}

def get_document_content(folder_name: str, file_name: str) -> dict:
    """Retrieve document content; supports PDF text extraction"""
    @retry_on_connection_error(max_retries=3, delay=1.0)
    def _download_file():
        file_path = _get_sp_path(f"{folder_name}/{file_name}")
        file = sp_context.web.get_file_by_server_relative_url(file_path)
        sp_context.load(file, ["Exists", "Length", "Name"])
        sp_context.execute_query()
        logger.info(f"File exists: {file.exists}, size: {file.length}")

        content = io.BytesIO()
        file.download(content)
        sp_context.execute_query()
        return content.getvalue(), file_name
    
    content_bytes, file_name = _download_file()
    
    # Determine file type and process accordingly
    lower_name = file_name.lower()
    file_type = next((t for t, exts in FILE_TYPES.items() if any(lower_name.endswith(ext) for ext in exts)), 'binary')
    
    if file_type == 'pdf':
        try:
            text, pages = extract_text_from_pdf(content_bytes)
            return {"name": file_name, "content_type": "text", "content": text, "original_type": "pdf", "page_count": pages, "size": len(content_bytes)}
        except Exception as e:
            logger.warning(f"PDF processing failed: {e}")
            return {"name": file_name, "content_type": "binary", "content_base64": base64.b64encode(content_bytes).decode(), "original_type": "pdf", "size": len(content_bytes)}
    
    if file_type == 'excel':
        try:
            text, sheets = extract_text_from_excel(content_bytes)
            return {"name": file_name, "content_type": "text", "content": text, "original_type": "excel", "sheet_count": sheets, "size": len(content_bytes)}
        except Exception as e:
            logger.warning(f"Excel processing failed: {e}")
            return {"name": file_name, "content_type": "binary", "content_base64": base64.b64encode(content_bytes).decode(), "original_type": "excel", "size": len(content_bytes)}
    
    if file_type == 'word':
        try:
            text, paragraphs = extract_text_from_word(content_bytes)
            return {"name": file_name, "content_type": "text", "content": text, "original_type": "word", "paragraph_count": paragraphs, "size": len(content_bytes)}
        except Exception as e:
            logger.warning(f"Word processing failed: {e}")
            return {"name": file_name, "content_type": "binary", "content_base64": base64.b64encode(content_bytes).decode(), "original_type": "word", "size": len(content_bytes)}
    
    if file_type == 'text':
        try:
            return {"name": file_name, "content_type": "text", "content": content_bytes.decode('utf-8'), "size": len(content_bytes)}
        except UnicodeDecodeError:
            pass
    
    return {"name": file_name, "content_type": "binary", "content_base64": base64.b64encode(content_bytes).decode(), "size": len(content_bytes)}

def download_document(folder_name: str, file_name: str, local_path: str) -> Dict[str, Any]:
    """Download document from SharePoint to local filesystem with fallback support"""
    logger.info(f"Downloading {folder_name}/{file_name} to {local_path}")
    
    try:
        # Get file from SharePoint
        file_path = _get_sp_path(f"{folder_name}/{file_name}")
        file = sp_context.web.get_file_by_server_relative_url(file_path)
        sp_context.load(file, ["Exists", "Length", "Name"])
        sp_context.execute_query()
        
        if not file.exists:
            return {"success": False, "error": f"File {file_name} does not exist in folder {folder_name}"}
        
        # Download file content
        content = io.BytesIO()
        file.download(content)
        sp_context.execute_query()
        content_bytes = content.getvalue()
        
        # Try to save to requested path first
        save_result = _save_content_to_file(content_bytes, local_path)
        
        if save_result["success"]:
            return {
                "success": True,
                "path": save_result["path"],
                "size": save_result["size"],
                "method": "primary"
            }
        
        # Fallback: save to fallback directory
        logger.warning(f"Primary path failed: {save_result['error']}, trying fallback")
        fallback_path = _get_fallback_path(file_name)
        fallback_result = _save_content_to_file(content_bytes, fallback_path)
        
        if fallback_result["success"]:
            return {
                "success": True,
                "path": fallback_result["path"], 
                "size": fallback_result["size"],
                "method": "fallback",
                "primary_error": save_result["error"]
            }
        
        # Both paths failed
        return {
            "success": False,
            "error": f"Both primary and fallback paths failed",
            "primary_error": save_result["error"],
            "fallback_error": fallback_result["error"]
        }
        
    except Exception as e:
        logger.error(f"Download failed for {folder_name}/{file_name}: {e}")
        return {"success": False, "error": f"Download operation failed: {str(e)}"}
